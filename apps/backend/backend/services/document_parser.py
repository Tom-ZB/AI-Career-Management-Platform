"""
Document Parser Service - Extract text from various document formats.
Supports PDF, DOCX, and TXT files.
"""
from typing import Optional
import io


class DocumentParser:
    """Service for parsing documents and extracting text content."""

    @staticmethod
    def parse_pdf(file_bytes: bytes) -> str:
        """
        Extract text from a PDF file using PyMuPDF.

        Args:
            file_bytes: PDF file content as bytes

        Returns:
            Extracted text content
        """
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text_parts = []

            for page in doc:
                text = page.get_text()
                text_parts.append(text)

            doc.close()
            return "\n".join(text_parts)

        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def parse_docx(file_bytes: bytes) -> str:
        """
        Extract text from a DOCX file using python-docx.

        Args:
            file_bytes: DOCX file content as bytes

        Returns:
            Extracted text content
        """
        try:
            from docx import Document

            doc = Document(io.BytesIO(file_bytes))
            text_parts = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n".join(text_parts)

        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    @staticmethod
    def parse_txt(file_bytes: bytes) -> str:
        """
        Extract text from a plain text file.

        Args:
            file_bytes: TXT file content as bytes

        Returns:
            Text content
        """
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ["latin-1", "cp1252", "gbk"]:
                try:
                    return file_bytes.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError("Failed to decode text file with any known encoding")

    @classmethod
    def parse(
        cls, file_bytes: bytes, content_type: Optional[str] = None, file_name: Optional[str] = None
    ) -> str:
        """
        Parse a document based on its content type or file extension.

        Args:
            file_bytes: File content as bytes
            content_type: MIME type of the file (optional)
            file_name: Original file name (optional, used to infer type)

        Returns:
            Extracted text content
        """
        # Determine file type from content type or file name
        if content_type:
            content_type = content_type.lower()
            if "pdf" in content_type:
                return cls.parse_pdf(file_bytes)
            elif "docx" in content_type or "word" in content_type:
                return cls.parse_docx(file_bytes)
            elif "text" in content_type or "plain" in content_type:
                return cls.parse_txt(file_bytes)

        # Fall back to file extension
        if file_name:
            ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
            if ext == "pdf":
                return cls.parse_pdf(file_bytes)
            elif ext == "docx":
                return cls.parse_docx(file_bytes)
            elif ext == "txt":
                return cls.parse_txt(file_bytes)

        # Try to detect by magic bytes
        if file_bytes[:4] == b"%PDF":
            return cls.parse_pdf(file_bytes)
        elif file_bytes[:8] == b"\x50\x4b\x03\x04\x14\x00\x06\x00":  # DOCX starts with PK
            return cls.parse_docx(file_bytes)

        # Default to text
        return cls.parse_txt(file_bytes)


def extract_text_from_document(file_bytes: bytes, content_type: Optional[str] = None, file_name: Optional[str] = None) -> str:
    """
    Convenience function to extract text from a document.

    Args:
        file_bytes: File content as bytes
        content_type: MIME type of the file (optional)
        file_name: Original file name (optional)

    Returns:
        Extracted text content
    """
    parser = DocumentParser()
    return parser.parse(file_bytes, content_type, file_name)
